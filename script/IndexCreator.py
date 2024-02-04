import pandas as pd
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
import re
import argparse


def create_index(file_name, sort_by='Topic', title="", ignoreCase=True, ignoreSymbol=True, groupNumbers=True):
    # Read the Excel file
    df = pd.read_excel(file_name)

    # Make all empty space an empty string
    df = df.fillna('')

    # Make it all strings
    df = df.astype(str)

    # Group by 'Topic' and join the 'Book' and 'Page' columns with a newline
    df = df.groupby(sort_by).agg({'Book': '\n'.join, 'Page': '\n'.join, 'Notes': ' '.join}).reset_index()

    if ignoreCase:
        if ignoreSymbol:
            # Sort the DataFrame by 'Topic' (case-insensitive and ignoring non-alphanumeric characters)
            df = df.sort_values(sort_by, key=lambda col: col.str.lower().map(lambda x: re.sub(r'\W+', '', x)))
        else:
            # Sort the DataFrame by 'Topic' (case-insensitive)
            df = df.sort_values(sort_by, key=lambda col: col.str.lower())
    elif ignoreSymbol:
        # Sort the DataFrame by 'Topic' (ignoring non-alphanumeric characters)
        df = df.sort_values(sort_by, key=lambda col: col.str.map(lambda x: re.sub(r'\W+', '', x)))
    else:
        # Sort the DataFrame by 'Topic'
        df = df.sort_values(sort_by)

    # Create a new Word document
    doc = Document()

    # Change the orientation of the document to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    if title:
        heading_title = doc.add_heading(title, 0)
        run = heading_title.runs[0]
        run.font.size = Pt(30)  # Set the font size to 30
        heading_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Initialize the current section
    current_section = ''
    table = None

    bShade = False
    # Iterate over the rows of the DataFrame
    for index, row in df.iterrows():
        # Get the first letter of the 'Topic'
        section = next((char for char in row[sort_by] if (char.isalpha() or char.isnumeric())), '')
        if ignoreCase:
            section = section.upper()
        if groupNumbers and section.isnumeric():
            section = '#'

        # If this is a new section, add a section header to the document
        if section and section != current_section:
            bShade = False
            section_header = doc.add_heading(section, level=1)
            run = section_header.runs[0]
            run.font.size = Pt(20)  # Set the font size to 20
            section_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            current_section = section

            # Add a table to the document
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Add the headers to the table
            for i, header in enumerate(['Topic', 'Book', 'Page', 'Notes']):
                cell = table.cell(0, i)
                cell.text = header
                paragraph = cell.paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

        # Add a new row to the table
        cells = table.add_row().cells

        if bShade:
            # Set the shading for even rows
            shading_elm = []
            for i in range(4):
                shading_elm.append(parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
                cells[i]._tc.get_or_add_tcPr().append(shading_elm[i])

        cells[0].text = f"{row['Topic']}"
        cells[1].text = f"{row['Book']}"
        cells[2].text = f"{row['Page']}"
        cells[3].text = "" if pd.isna(row['Notes']) else f"{row['Notes']}"

        bShade = not bShade

    # Save the Word document
    doc.save('../docs/index.docx')


# Call the function with the name of your Excel file
if __name__ == '__main__':
    # Create the parser
    parser = argparse.ArgumentParser(description="This is a tool used to help create an index for any resource. It requires an xlsx file as input and outputs to ../docs/index.docx.")

    # Add the arguments
    parser.add_argument('-f', metavar='filename', type=str, help='Path to index xlsx file.', required=True)
    parser.add_argument('--recommended', action='store_true', help='Use recommended settings.')
    parser.add_argument('-s', metavar='sort', type=str, help='Topic to sort by.', default='')
    parser.add_argument('-t', metavar='title', nargs='+', type=str, help='Title for the page.', default='')
    parser.add_argument('-Ic', action='store_true', help='Ignore case when sorting.')
    parser.add_argument('-Is', action='store_true', help='Ignore symbols when sorting.')
    parser.add_argument('-g', action='store_true', help='Ignore numbers')

    # Execute the parse_args() method
    args = parser.parse_args()
    title = ' '.join(args.t)
    print(f"Creating index based on {args.f} in ../docs/index.docx...")

    if args.recommended:
        sortColumn = 'Topic'

        if args.s != '':
            sortColumn = args.s

        create_index(args.f, sortColumn, title, True, True, True)
    else:
        create_index(args.f, args.s, title, args.Ic, args.Is, args.g)

    print("Done.")